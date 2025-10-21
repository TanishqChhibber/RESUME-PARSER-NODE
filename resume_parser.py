import sys
import json
import os
import fitz  # PyMuPDF for extracting links
import pdfplumber
import requests
import yaml
import re
from docx import Document
from PIL import Image
import pytesseract
import io

CONFIG_PATH = "config.yaml"
SUPPORTED_FORMATS = {
    'pdf': ['.pdf'],
    'word': ['.doc', '.docx'],
    'image': ['.png', '.jpg', '.jpeg'],
    'text': ['.txt']
}
EXTRACTION_TEMPLATE = """Extract the following ATS fields as JSON:
{
  "name": "Full name",
  "email": "Email address",
  "linkedin": "LinkedIn profile URL",
  "phone": "Phone number",
  "github": "GitHub profile URL",
  "skills": ["list", "of", "skills"],
  "experience": [
    {
      "company": "Company name",
      "role": "Job title",
      "duration": "From–To",
      "details": ["detailed bullet points of responsibilities/achievements"]
    }
  ],
  "projects": [
    {
      "title": "Project name",
      "duration": "From–To",
      "details": ["key points describing the project"]
    }
  ],
  "education": ["Degree details"]
}"""


def load_api_key():
    try:
        with open(CONFIG_PATH) as file:
            data = yaml.safe_load(file)
            api_key = data.get('OPENROUTER_API_KEY') or data.get('openrouter')
            if not api_key:
                raise ValueError("Missing API key in config.yaml")
            return api_key
    except Exception as e:
        print(json.dumps({"error": f"Config error: {str(e)}"}))
        sys.exit(1)

def get_file_type(file_path):
    """Determine file type from extension"""
    ext = os.path.splitext(file_path)[1].lower()
    for file_type, extensions in SUPPORTED_FORMATS.items():
        if ext in extensions:
            return file_type
    return None

def extract_text_from_pdf(file_path):
    """Extract text from PDF using pdfplumber"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text(x_tolerance=1, y_tolerance=1)
                if extracted_text:
                    text += extracted_text + "\n\n"
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")
    return text

def extract_text_from_word(file_path):
    """Extract text from Word document"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        raise Exception(f"Word document extraction error: {str(e)}")
    return text

def extract_text_from_image(file_path):
    """Extract text from image using OCR"""
    try:
        # Open image using PIL
        image = Image.open(file_path)
        
        # Convert to RGB if necessary (for PNG transparency)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        if not text.strip():
            raise Exception("No text was extracted from the image")
            
    except Exception as e:
        raise Exception(f"Image extraction error: {str(e)}")
    return text

def extract_text_from_txt(file_path):
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try different encodings if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Text file reading error: {str(e)}")
    except Exception as e:
        raise Exception(f"Text file reading error: {str(e)}")

def extract_text(file_path):
    """Extract text from various file formats"""
    if not os.path.exists(file_path):
        print(json.dumps({"error": "File not found"}))
        sys.exit(1)

    try:
        # Determine file type
        file_type = get_file_type(file_path)
        if not file_type:
            print(json.dumps({"error": f"Unsupported file format. Supported formats: {', '.join([ext for exts in SUPPORTED_FORMATS.values() for ext in exts])}"}))
            sys.exit(1)

        # Extract text based on file type
        if file_type == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif file_type == 'word':
            text = extract_text_from_word(file_path)
        elif file_type == 'image':
            text = extract_text_from_image(file_path)
        elif file_type == 'text':
            text = extract_text_from_txt(file_path)
        else:
            print(json.dumps({"error": "Unsupported file format"}))
            sys.exit(1)

        if not text.strip():
            print(json.dumps({"error": "No text could be extracted from the file"}))
            sys.exit(1)

        return text.strip()

    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


def extract_links_from_pdf(file_path):
    """Extract links from PDF file"""
    github_links = []
    linkedin_links = []
    try:
        doc = fitz.open(file_path)
        for page in doc:
            links = page.get_links()
            for link in links:
                url = link.get("uri", "")
                if "github.com" in url:
                    github_links.append(url)
                elif "linkedin.com/in" in url or "linkedin.com/" in url:
                    linkedin_links.append(url)
    except Exception as e:
        print(f"Warning: Failed to extract links from PDF: {str(e)}", file=sys.stderr)
    return github_links, linkedin_links

def extract_links_from_text(text):
    """Extract links from raw text using regex"""
    github_links = re.findall(r'https?://(?:www\.)?github\.com/[A-Za-z0-9_-]+/?(?!\S)', text)
    linkedin_links = re.findall(r'https?://(?:www\.)?linkedin\.com/(?:in|profile)/[A-Za-z0-9_-]+/?(?!\S)', text)
    return github_links, linkedin_links

def extract_links(file_path):
    """Extract links from various file formats"""
    github_links = []
    linkedin_links = []
    
    try:
        # Get file type
        file_type = get_file_type(file_path)
        
        # Extract links based on file type
        if file_type == 'pdf':
            # Use PDF-specific extraction
            github_links, linkedin_links = extract_links_from_pdf(file_path)
        
        # For all file types, also try to extract from raw text
        # This serves as a fallback and can find plain text URLs
        text = extract_text(file_path)
        text_github_links, text_linkedin_links = extract_links_from_text(text)
        
        # Combine and deduplicate links
        github_links = list(set(github_links + text_github_links))
        linkedin_links = list(set(linkedin_links + text_linkedin_links))
    
    except Exception as e:
        print(json.dumps({"error": f"Failed to extract links: {str(e)}"}))
    
    return github_links, linkedin_links

def extract_basic_info(text):
    """Extract basic information using regex patterns"""
    
    # Extract name (usually first line or early in document)
    lines = text.split('\n')
    name = "Not found"
    
    for i, line in enumerate(lines[:10]):  # Check first 10 lines
        line = line.strip()
        if not line or line.lower() in ['resume', 'cv', 'curriculum vitae']:
            continue
        
        # Check if line looks like a name (2-4 words, mostly alphabetic, title case)
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(word.replace('.', '').replace(',', '').isalpha() and word[0].isupper() for word in words):
                name = line
                break
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_matches = re.findall(email_pattern, text)
    email = email_matches[0] if email_matches else "Not found"
    
    # Extract phone
    phone_pattern = r'(\(\+\d{1,3}\)\s?\d{10}|\+\d{1,3}[\s\-]?\d{10}|\(\+\d{1,3}\)\d{10}|\d{10})'
    phone_matches = re.findall(phone_pattern, text)
    phone = phone_matches[0] if phone_matches else "Not found"
    
    return {"name": name, "email": email, "phone": phone}

def extract_skills(text):
    """Extract skills from resume text"""
    skills = []
    
    # Look for skills section
    skills_section_match = re.search(r'SKILLS\s+(.*?)(?:PROJECTS|CERTIFICATIONS|ACHIEVEMENTS|EXPERIENCE|EDUCATION|\Z)', text, re.DOTALL | re.IGNORECASE)
    
    if skills_section_match:
        skills_text = skills_section_match.group(1).strip()
        lines = skills_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if '|' in line:
                # Format: "Category Item1 | Item2 | Item3"
                parts = line.split(None, 1)  # Split on first whitespace
                if len(parts) >= 2:
                    items_text = parts[1]
                    items = [item.strip() for item in items_text.split('|') if item.strip()]
                    skills.extend(items)
    
    # Fallback: common technical skills
    if not skills:
        skill_patterns = [
            r'\b(?:Python|JavaScript|Java|C\+\+|SQL|HTML|CSS|React|Node\.js)\b',
            r'\b(?:MySQL|MongoDB|PostgreSQL|Redis|AWS|Docker|Git)\b',
            r'\b(?:TensorFlow|PyTorch|Pandas|NumPy|OpenCV|Matplotlib)\b',
            r'\b(?:Power BI|Tableau|Excel|Visual Studio|VS Code)\b'
        ]
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.extend(matches)
    
    # Clean and deduplicate
    cleaned_skills = []
    for skill in skills:
        skill = skill.strip()
        # Filter out dates, common false positives, and non-skills
        if (skill and 
            len(skill) > 1 and 
            len(skill) < 30 and 
            skill not in cleaned_skills and
            not re.search(r'\b\d{4}\b', skill) and  # Filter out years/dates
            not re.search(r'(?:January|February|March|April|May|June|July|August|September|October|November|December)', skill, re.IGNORECASE) and
            not skill.lower() in ['experience', 'projects', 'education', 'skills', 'languages', 'databses', 'libraries', 'frameworks']):
            cleaned_skills.append(skill)
    
    return cleaned_skills[:15]

def extract_experience(text):
    """Extract work experience with improved parsing for complex formats"""
    experiences = []
    
    # Look for experience section
    exp_section_match = re.search(r'EXPERIENCE\s+(.*?)(?:PROJECTS|SKILLS|EDUCATION|CERTIFICATIONS|\Z)', text, re.DOTALL | re.IGNORECASE)
    
    if exp_section_match:
        exp_text = exp_section_match.group(1).strip()
        
        # Check for different experience formats
        lines = [line.strip() for line in exp_text.split('\n') if line.strip()]
        
        # Check if this is a skills/expertise format
        # Example: "in AI-driven solutions startup leadership, and data"
        if any('in' in line.lower() or 'and' in line.lower() for line in lines[:2]):
            # First, check if there's a structured internship/job entry
            company_line_idx = -1
            for i, line in enumerate(lines):
                if ('Technologies' in line or 'Tech' in line) and ('Intern' in line or 'Internship' in line):
                    company_line_idx = i
                    break
            
            if company_line_idx >= 0:
                # This is a structured experience entry
                company_line = lines[company_line_idx]
                company = company_line.replace('Summer Internship', '').replace('Internship', '').strip()
                if company_line.endswith('Summer Internship'):
                    duration = "Summer Internship"
                elif company_line.endswith('Internship'):
                    duration = "Internship"
                else:
                    duration = "Current"
                
                # Look for role in next line
                role = "Software Development Intern"  # Default
                if company_line_idx + 1 < len(lines):
                    next_line = lines[company_line_idx + 1].strip()
                    if not next_line.startswith(('•', '–', '-')) and 'Intern' in next_line:
                        role = next_line
                
                # Collect all bullet points
                details = []
                current_point = ""
                
                for line in lines[company_line_idx + 1:]:
                    line = line.strip()
                    if not line:
                        continue
                        
                    if line.startswith(('•', '–', '-')):
                        if current_point:
                            details.append(current_point)
                        current_point = line[1:].strip()
                    else:
                        if current_point:
                            current_point += " " + line
                        elif not line.startswith(('•', '–', '-')) and 'Intern' not in line:
                            current_point = line
                
                if current_point:
                    details.append(current_point)
                
                # Filter and clean details
                cleaned_details = []
                for detail in details:
                    # Clean up the detail text
                    detail = re.sub(r'\s+', ' ', detail).strip()
                    # Only keep substantial details
                    if len(detail) > 15 and not detail.endswith(('and', 'or', 'for', 'with', 'to')):
                        cleaned_details.append(detail)
                
                experiences.append({
                    "company": company,
                    "role": role,
                    "duration": duration,
                    "details": cleaned_details[:5]  # Limit to 5 most important points
                })
            
            # Also check for descriptive experience text
            expertise_lines = []
            for line in lines:
                if 'in' in line.lower() and any(keyword in line.lower() for keyword in ['ai', 'data', 'solution', 'startup', 'leadership']):
                    # Split on common delimiters but preserve meaningful phrases
                    expertise_line = line.strip()
                    expertise_line = re.sub(r'\s+', ' ', expertise_line)
                    expertise_lines.append(expertise_line)
            
            if expertise_lines:
                experiences.append({
                    "company": "Professional Experience",
                    "role": "Domain Expertise",
                    "duration": "Current",
                    "details": expertise_lines
                })
            
            return experiences
        
        # Format 1: Bullet point format with company info
        # • Mathrithms Technologies Pvt. Ltd. (Aikin) Summer Internship
        # Software Development Intern
        if lines and lines[0].startswith('•'):
            # Parse bullet point format
            i = 0
            while i < len(lines):
                line = lines[i]
                if line.startswith('•'):
                    # Extract company info from bullet line
                    company_line = line[1:].strip()
                    
                    # Extract company name and internship type
                    if 'Internship' in company_line or 'Intern' in company_line:
                        # Split by internship keyword
                        if 'Summer Internship' in company_line:
                            company = company_line.replace('Summer Internship', '').strip()
                            duration = "Summer Internship"
                        elif 'Internship' in company_line:
                            company = company_line.replace('Internship', '').strip()
                            duration = "Internship"
                        else:
                            company = company_line
                            duration = "Unknown Duration"
                    else:
                        company = company_line
                        duration = "Unknown Duration"
                    
                    # Next line should be the role
                    role = "Unknown Role"
                    if i + 1 < len(lines) and not lines[i + 1].startswith('•') and not lines[i + 1].startswith('–'):
                        role = lines[i + 1].strip()
                        i += 1
                    
                    # Collect responsibilities (lines starting with –)
                    responsibilities = []
                    i += 1
                    while i < len(lines) and (lines[i].startswith('–') or lines[i].startswith('-')):
                        resp = lines[i][1:].strip()
                        if resp and len(resp) > 15:
                            responsibilities.append(resp)
                        i += 1
                    
                    if company and role:
                        experiences.append({
                            "company": company,
                            "role": role,
                            "duration": duration,
                            "details": responsibilities[:5]
                        })
                    
                    continue
                i += 1
        
        # Format 2: Descriptive experience (fallback for unstructured)
        elif not any(',' in line or '|' in line for line in lines[:3]):
            # This appears to be a descriptive experience section
            description_lines = []
            for line in lines:
                if len(line) > 10 and not line.lower().startswith(('experience', 'work', 'employment')):
                    description_lines.append(line)
            
            if description_lines:
                experiences.append({
                    "company": "Professional Experience",
                    "role": "Various Roles", 
                    "duration": "Ongoing",
                    "details": description_lines[:3]
                })
        
        # Format 3: Structured format (original logic)
        else:
            # Split by company entries - look for patterns like "Company, Role | Date" or "Company, Role Date"
            company_pattern = r'(?=(?:SwitchiT|Celebal Technologies|CipherSquare Technologies|[A-Z][a-zA-Z\s]+ Technologies|[A-Z][a-zA-Z\s]+(?:Corp|Inc|LLC|Ltd|Company|Systems|Solutions|Group)))'
            company_entries = re.split(company_pattern, exp_text)
            
            for entry in company_entries:
                entry = entry.strip()
                if len(entry) > 50:  # Filter out short entries
                    lines = [line.strip() for line in entry.split('\n') if line.strip()]
                    if not lines:
                        continue
                        
                    first_line = lines[0]
                    
                    if ',' in first_line:
                        parts = first_line.split(',', 1)
                        company = parts[0].strip()
                        rest = parts[1].strip()
                        
                        # Extract role and duration
                        if '|' in rest:
                            role_part, duration_part = rest.split('|', 1)
                            role = role_part.strip()
                            duration = duration_part.strip()
                        else:
                            # Format without pipe - extract date pattern from end
                            date_match = re.search(r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}(?:\s*[–\-]\s*(?:Present|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}))?)$', rest)
                            if date_match:
                                duration = date_match.group(1)
                                role = rest[:date_match.start()].strip()
                            else:
                                role = rest
                                duration = "Unknown Duration"
                    else:
                        company = first_line.split()[0] if first_line.split() else "Unknown Company"
                        role = "Unknown Role"
                        duration = "Unknown Duration"
                    
                    # Extract responsibilities (bullet points)
                    responsibilities = []
                    for line in lines[1:]:
                        if line.startswith('•') or line.startswith('-'):
                            resp = line[1:].strip()
                            if resp and len(resp) > 10:
                                responsibilities.append(resp)
                    
                    if company and role:
                        experiences.append({
                            "company": company,
                            "role": role,
                            "duration": duration,
                            "details": responsibilities[:5]
                        })
    
    return experiences

def extract_projects(text):
    """Extract projects from resume with improved parsing for complex formats"""
    projects = []
    
    # Look for projects section with variations of "PROJECTS" or "KEY PROJECTS"
    projects_section_match = re.search(r'(?:KEY\s+)?PROJECTS\s+(.*?)(?:RELEVANTCOURSEWORK|RELEVANT COURSEWORK|KEY COURSES|EXTRACURRICULAR|EDUCATION|SKILLS|CERTIFICATIONS|ACHIEVEMENTS|TECHNICAL SKILLS|\Z)', text, re.DOTALL | re.IGNORECASE)
    
    if projects_section_match:
        projects_text = projects_section_match.group(1).strip()
        
        # Split by project entries that start with "–" or "-" followed by a title
        project_entries = re.split(r'(?m)^[–\-]\s+', projects_text)
        # Filter out empty entries
        project_entries = [entry.strip() for entry in project_entries if entry.strip()]
        
        for entry in project_entries:
            # Skip empty entries
            if not entry:
                continue
                
            lines = [line.strip() for line in entry.split('\n') if line.strip()]
            if not lines:
                continue
            
            # First line should contain project title and duration
            first_line = lines[0]
            
            # Try to extract duration if it exists
            duration_match = re.search(r'(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[.\s]+\d{4}(?:\s*[-–]\s*(?:Present|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[.\s]+\d{4}))?)', first_line)
            
            if duration_match:
                duration = duration_match.group(0)
                title = first_line[:duration_match.start()].strip()
            else:
                title = first_line
                duration = "Not specified"
            
            # Clean up title
            title = re.sub(r'\s*[–\-]\s*$', '', title)  # Remove trailing dashes
            
            # Extract development type/organization from second line if it exists
            org = ""
            if len(lines) > 1 and not any(lines[1].startswith(x) for x in ('∗', '•', '-', '–', '*')):
                org = lines[1].strip()
            
            # Collect project details
            details = []
            current_detail = ""
            
            for line in lines[1:]:
                # Skip organization line
                if line == org:
                    continue
                    
                # Look for lines starting with bullet points
                if any(line.startswith(x) for x in ('∗', '•', '-', '–', '*')):
                    if current_detail:  # Save previous detail if exists
                        if len(current_detail) > 15 and not any(current_detail.lower().startswith(x) for x in ('tool', 'project link', 'technologies used')):
                            details.append(current_detail)
                    current_detail = line[1:].strip()  # Start new detail
                else:
                    # Continue previous detail if it exists
                    if current_detail:
                        current_detail += " " + line.strip()
                        
            # Don't forget the last detail
            if current_detail and len(current_detail) > 15 and not any(current_detail.lower().startswith(x) for x in ('tool', 'project link', 'technologies used')):
                details.append(current_detail)
            
            # Clean up details
            cleaned_details = []
            for detail in details:
                # Remove common prefixes
                detail = re.sub(r'^[∗•\-–\*]\s*', '', detail)
                # Clean up whitespace
                detail = re.sub(r'\s+', ' ', detail).strip()
                if detail and len(detail) > 15 and not any(detail.lower().startswith(x) for x in ('tool', 'project link', 'technologies used')):
                    cleaned_details.append(detail)
            
            # Only add projects with meaningful titles
            if title and len(title) > 3 and not any(x.lower() in title.lower() for x in ['individual project', 'development project']):
                project_data = {
                    "title": title,
                    "duration": duration,
                    "details": cleaned_details[:4]  # Limit to 4 most important details
                }
                # Add organization if it exists and is meaningful
                if org and not any(x.lower() in org.lower() for x in ['tool', 'project link', 'technologies used', 'individual project', 'development project']):
                    project_data["organization"] = org
                    
                projects.append(project_data)
    
    return projects

def extract_education(text):
    """Extract education information"""
    education = []
    
    # Look for education section
    edu_section_match = re.search(r'EDUCATION\s+(.*?)(?:EXPERIENCE|PROJECTS|SKILLS|CERTIFICATIONS|\Z)', text, re.DOTALL | re.IGNORECASE)
    
    if edu_section_match:
        edu_text = edu_section_match.group(1).strip()
        lines = edu_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 20:  # Filter out short lines
                education.append(line)
    
    return education

def get_github_profile(github_links):
    """Filters GitHub profile link by removing project links."""
    for link in github_links:
        match = re.match(r"https?://github\.com/([^/]+)$", link)
        if match:
            return link  # Return first valid profile link
    return None  # No valid GitHub profile found

def ats_extractor(resume_data):
    """
    Extracts structured data from resume text using GPT API.
    Includes retry logic and enhanced error handling.
    """
    try:
        api_key = load_api_key()

        # Retry logic with exponential backoff
        retries = 3
        for attempt in range(retries):
            try:
                response = requests.post(
                    url="https://api.openai.com/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "gpt-4",
                        "messages": [
                            {"role": "system", "content": "You are an ATS parser. Return ONLY valid JSON. Do not include explanations."},
                            {"role": "user", "content": f"{EXTRACTION_TEMPLATE}\n\nResume data:\n{resume_data}"}
                        ],
                        "temperature": 0.1,
                        "max_tokens": 4000
                    },
                    timeout=60  # Increased timeout duration
                )

                response.raise_for_status()
                content = response.json()["choices"][0]["message"]["content"]
                cleaned = content.strip().replace('```json', '').replace('```', '')

                return json.loads(cleaned)

            except requests.exceptions.Timeout:
                if attempt < retries - 1:
                    continue  # Retry on timeout
                else:
                    raise  # Raise the last timeout exception

    except json.JSONDecodeError:
        return {"error": "API returned invalid JSON"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Request failed: {str(e)}"}
    except Exception as e:
        return {"error": f"Processing failed: {str(e)}"}


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file provided"}))
        sys.exit(1)

    file_path = sys.argv[1]
    extracted_text = extract_text(file_path)

    if not extracted_text:
        print(json.dumps({"error": "No text extracted"}))
        sys.exit(1)

    # Extract links
    github_links, linkedin_links = extract_links(file_path)
    github_profile = get_github_profile(github_links)

    # Try AI extraction first
    parsed_data = ats_extractor(extracted_text)

    # If AI extraction failed or returned error, use fallback extraction
    if isinstance(parsed_data, dict) and "error" in parsed_data:
        print(f"AI extraction failed: {parsed_data['error']}", file=sys.stderr)

        # Use fallback extraction methods
        basic_info = extract_basic_info(extracted_text)
        skills = extract_skills(extracted_text)
        experience = extract_experience(extracted_text)
        projects = extract_projects(extracted_text)
        education = extract_education(extracted_text)

        parsed_data = {
            "name": basic_info["name"],
            "email": basic_info["email"],
            "phone": basic_info["phone"],
            "linkedin": linkedin_links[0] if linkedin_links else "Not found",
            "github": github_profile or "Not found",
            "skills": skills,
            "experience": experience,
            "projects": projects,
            "education": education,
            "extraction_method": "fallback"
        }
    else:
        # AI extraction succeeded, enhance with links
        parsed_data["github"] = github_profile or "Not found"
        parsed_data["linkedin"] = linkedin_links[0] if linkedin_links else "Not found"
        parsed_data["extraction_method"] = "ai"

    print(json.dumps(parsed_data, indent=2))
